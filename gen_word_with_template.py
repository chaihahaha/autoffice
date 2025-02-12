import re
import docx
from docx import Document
from docx.oxml.ns import qn
import os
import csv

master_doc = Document('master_template.docx')
empty_table_doc = Document('empty_table.docx')
N = 1

# Function to append a template's content into the master document
def append_template_content(src_doc, dest_doc, run):
    pict_element = None
    for element in src_doc.element.body:
        if element.tag.endswith('p'):
            for child in element:
                if child.tag.endswith('r'):
                    for sub_child in child:
                        if sub_child.tag.endswith('pict'):
                            pict_element = sub_child
                            break

    run._r.append(pict_element)
    return

def replace_text_in_docx(doc_path, entry_pairs, patterns):
    # Load the template document
    doc = Document(doc_path)
    #print(doc.element.body.xml)
    #exit()

    patterns = [re.compile(i) for i in patterns]

    # Function to replace text in a given XML element
    def replace_text_in_element(element, entry_list):
        for child in element.iter():
            if child.tag == qn('w:t'):  # A text element
                for entry, p in zip(entry_list, patterns):
                    if p.search(child.text):
                        child.text = p.sub(f"{entry}", child.text)

    paragraph = master_doc.add_paragraph()
    run = paragraph.add_run()

    # Iterate over the entry pairs to create a modified document for each pair
    for i in range(len(entry_pairs)):
        entry = entry_pairs[i]
        if entry == "newpage":
            paragraph = master_doc.add_paragraph()
            run = paragraph.add_run()
            master_doc.add_page_break()
            paragraph = master_doc.add_paragraph()
            run = paragraph.add_run()
            run.add_tab()
            #print(entry_pairs[i-1])
        else:
            entry_list = list(entry)
            # Create a copy of the loaded document
            new_doc = Document(doc_path)

            # Replace text in the main document body
            replace_text_in_element(new_doc.element.body, entry_list)

            global N
            # Save the modified document with a unique name
            N += 1
            #if i+1 < len(entry_pairs) and entry_pairs[i+1] == "newpage":
            #    run.add_tab()
            #    run.add_tab()
            append_template_content(new_doc, master_doc, run)
            run.add_tab()
            print(f"items: {N}")

